from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

def actualizar_documento_word():
    # Ruta del documento
    docx_path = r"C:\Users\Carmo\Downloads\Tabla de Contenido Proyecto GradoV3 FORMACERO.docx"

    if not os.path.exists(docx_path):
        print(f"Documento no encontrado: {docx_path}")
        return

    doc = Document(docx_path)

    print("Actualizando documento Word...")

    # 1. ACTUALIZAR INTRODUCCIÓN
    actualizar_introduccion(doc)

    # 2. AGREGAR DIAGRAMAS UML
    agregar_diagramas_uml(doc)

    # 3. COMPLETAR MODELO C4
    agregar_modelo_c4(doc)

    # 4. AGREGAR DESCRIPCIÓN DE PÁGINAS FRONTEND
    agregar_paginas_frontend(doc)

    # 5. COMPLETAR ESPECIFICACIÓN BD
    agregar_especificacion_bd(doc)

    # 6. AGREGAR CONCLUSIONES Y RECOMENDACIONES
    agregar_conclusiones_recomendaciones(doc)

    # 7. AGREGAR BIBLIOGRAFÍA
    agregar_bibliografia(doc)

    # 8. AGREGAR MANUAL DE USUARIO
    agregar_manual_usuario(doc)

    # 9. ACTUALIZAR ALCANCE (EXCLUIR NÓMINA Y VACACIONES)
    actualizar_alcance(doc)

    # Guardar documento actualizado
    nuevo_path = r"C:\Users\Carmo\Downloads\Tabla de Contenido Proyecto GradoV3 FORMACERO_ACTUALIZADO.docx"
    doc.save(nuevo_path)
    print(f"Documento actualizado guardado como: {nuevo_path}")

def actualizar_introduccion(doc):
    """Actualiza la sección de introducción"""
    # Buscar el párrafo de introducción
    for para in doc.paragraphs:
        if "Introducción" in para.text and "(realizar al completar el resto del documento)" in para.text:
            para.text = """Introducción

El presente documento describe el desarrollo e implementación de una plataforma web integral para la gestión de talento humano y automatización de certificados laborales en Formacero S.A.S. Esta solución tecnológica surge como respuesta a las necesidades identificadas en el proceso manual de gestión de información laboral de la empresa.

La plataforma desarrollada permite centralizar, gestionar y consultar la información de los empleados de manera digital, garantizando la integridad, seguridad y confidencialidad de los datos conforme a la Ley 1581 de Protección de Datos Personales en Colombia. La implementación incluye módulos de registro de hoja de vida, gestión documental, autogestión del empleado y generación automática de certificados laborales.

El sistema se ha desarrollado siguiendo estándares de calidad en ingeniería de software, utilizando arquitecturas modernas y tecnologías web actuales, lo que garantiza escalabilidad, mantenibilidad y una experiencia de usuario óptima tanto en dispositivos móviles como de escritorio."""

def agregar_diagramas_uml(doc):
    """Agrega todos los diagramas UML"""
    # Agregar sección de diagramas UML
    doc.add_heading('Diagramas UML', level=1)

    # Diagrama de Casos de Uso
    doc.add_heading('Diagrama de Casos de Uso', level=2)
    doc.add_paragraph("""
El diagrama de casos de uso muestra las interacciones entre los actores del sistema (Administrador y Empleado) y las funcionalidades principales:

- **Administrador**: Gestiona empleados, reportes, certificados, búsqueda avanzada, auditoría
- **Empleado**: Consulta perfil, edita datos limitados, gestiona documentos, genera certificados
- **Sistema**: Maneja autenticación, logout automático, recuperación de contraseña

Los casos de uso están detallados en las especificaciones correspondientes.""")
    # Aquí iría la imagen del diagrama

    # Diagrama de Clases
    doc.add_heading('Diagrama de Clases', level=2)
    doc.add_paragraph("""
El diagrama de clases representa la estructura del modelo de datos y las relaciones entre las entidades principales:

**Entidades principales:**
- Usuario: Gestión de autenticación y roles
- Empleado: Información laboral y personal
- ExEmpleado: Historial de empleados retirados
- Reporte: Reportes administrativos
- Certificado: Documentos generados
- DocumentoEmpleado: Archivos asociados
- Departamento: Estructura organizacional

**Controladores:**
- AuthController: Manejo de autenticación
- EmpleadosController: CRUD de empleados
- ReportesController: Gestión de reportes

**Servicios:**
- ApiService: Comunicación HTTP
- EmpleadosService: Lógica de negocio de empleados""")

    # Diagrama de Componentes
    doc.add_heading('Diagrama de Componentes', level=2)
    doc.add_paragraph("""
La arquitectura de componentes se divide en tres capas principales:

**Frontend (React + Vite):**
- Componentes de interfaz de usuario
- Servicios de API
- Utilidades de autenticación

**Backend (Node.js + Express):**
- Rutas de API REST
- Controladores de lógica de negocio
- Middlewares de autenticación y autorización

**Base de Datos (Supabase):**
- PostgreSQL para datos relacionales
- Storage para archivos multimedia""")

    # Diagrama de Actividades
    doc.add_heading('Diagrama de Actividades', level=2)
    doc.add_paragraph("""
El flujo de actividades del sistema incluye:

1. **Autenticación**: Login con validación JWT
2. **Gestión de Empleados**: CRUD completo con validaciones
3. **Gestión Documental**: Upload/download con restricciones
4. **Generación de Certificados**: Creación automática de PDF
5. **Auditoría**: Registro de cambios sensibles
6. **Búsqueda**: Filtrado avanzado de empleados""")

    # Diagrama de Despliegue
    doc.add_heading('Diagrama de Despliegue', level=2)
    doc.add_paragraph("""
La infraestructura de despliegue incluye:

- **Cliente**: Navegadores modernos con soporte responsive
- **Servidor**: Node.js con Express en puerto 3001
- **Base de Datos**: Supabase PostgreSQL con backup automático
- **Almacenamiento**: Supabase Storage para archivos
- **Email**: SMTP para notificaciones""")

def agregar_modelo_c4(doc):
    """Agrega el modelo C4 completo"""
    doc.add_heading('Modelo C4 - Arquitectura del Sistema', level=1)

    doc.add_heading('C4-1: Diagrama de Contexto', level=2)
    doc.add_paragraph("""
El diagrama de contexto muestra el sistema en su entorno, identificando:

- **Usuarios externos**: Administrador RRHH y Empleados
- **Sistema principal**: Plataforma web de gestión RRHH
- **Sistemas externos**: Supabase (BD + Storage) y servicio de email SMTP

Este diagrama proporciona una visión de alto nivel de las interacciones del sistema con su entorno.""")

    doc.add_heading('C4-2: Diagrama de Contenedores', level=2)
    doc.add_paragraph("""
Los contenedores representan las aplicaciones tecnológicas principales:

- **SPA React + Vite**: Interfaz de usuario en el navegador
- **API Node.js + Express**: Servidor de aplicaciones backend
- **Supabase PostgreSQL**: Base de datos relacional
- **Supabase Storage**: Almacenamiento de archivos
- **SMTP Server**: Servicio de envío de correos

La comunicación se realiza mediante HTTP/HTTPS con autenticación JWT.""")

    doc.add_heading('C4-3: Diagrama de Componentes', level=2)
    doc.add_paragraph("""
Los componentes detallan la estructura interna de cada contenedor:

**Frontend:**
- App.jsx: Componente principal
- Router: Navegación React
- Pages: Componentes de vistas
- Services: Lógica de negocio del cliente

**Backend:**
- Server.js: Servidor HTTP
- Routes: Definición de endpoints
- Controllers: Lógica de negocio
- Middleware: Autenticación y autorización""")

    doc.add_heading('C4-4: Diagrama de Código', level=2)
    doc.add_paragraph("""
El diagrama de código muestra las clases principales y sus relaciones:

- **Controladores**: AuthController, EmpleadosController, ReportesController
- **Servicios**: ApiService, EmpleadosService
- **Middleware**: AuthMiddleware para protección de rutas
- **Cliente BD**: SupabaseClient para acceso a datos

Este nivel muestra la implementación concreta en código.""")

def agregar_paginas_frontend(doc):
    """Agrega descripción completa de las páginas del frontend"""
    doc.add_heading('Páginas del Sistema Frontend', level=2)

    paginas = [
        {
            "ruta": "/login",
            "nombre": "Inicio de Sesión",
            "descripcion": "Página de autenticación con campos usuario/correo y contraseña. Incluye enlace a recuperación de contraseña.",
            "roles": "Todos"
        },
        {
            "ruta": "/dashboard",
            "nombre": "Panel Principal",
            "descripcion": "Dashboard con menú de navegación, estadísticas de empleados, búsqueda avanzada y acceso rápido a funciones principales.",
            "roles": "Admin, Empleado"
        },
        {
            "ruta": "/registrar-empleados",
            "nombre": "Registro de Empleados",
            "descripcion": "Formulario completo para crear nuevos empleados con subida de foto y documentos. Validaciones de campos requeridos.",
            "roles": "Admin"
        },
        {
            "ruta": "/informacion-empleados",
            "nombre": "Gestión de Empleados",
            "descripcion": "Lista completa de empleados activos con opciones de editar, eliminar y ver detalles. Paginación incluida.",
            "roles": "Admin"
        },
        {
            "ruta": "/empleado-detalle/:id",
            "nombre": "Perfil de Empleado",
            "descripcion": "Vista detallada del perfil de un empleado con información laboral, documentos y opciones de edición según rol.",
            "roles": "Admin, Empleado"
        },
        {
            "ruta": "/organizacion",
            "nombre": "Estructura Organizacional",
            "descripcion": "Vista jerárquica de la empresa mostrando departamentos y empleados. Implementa la pirámide empresarial.",
            "roles": "Admin"
        },
        {
            "ruta": "/lista-exempleados",
            "nombre": "Historial Ex-Empleados",
            "descripcion": "Lista de empleados retirados con información histórica para auditoría legal.",
            "roles": "Admin"
        },
        {
            "ruta": "/certificado-laboral",
            "nombre": "Generador de Certificados",
            "descripcion": "Interfaz para generar y descargar certificados laborales en formato PDF.",
            "roles": "Admin, Empleado"
        },
        {
            "ruta": "/reportes",
            "nombre": "Gestión de Reportes",
            "descripcion": "Sistema de reportes administrativos con creación, consulta y gestión de reportes de empleados.",
            "roles": "Admin"
        },
        {
            "ruta": "/forgot-password",
            "nombre": "Recuperar Contraseña",
            "descripcion": "Formulario para solicitar recuperación de contraseña vía email con token temporal.",
            "roles": "Todos"
        },
        {
            "ruta": "/reset-password",
            "nombre": "Restablecer Contraseña",
            "descripcion": "Página para ingresar nueva contraseña usando token de recuperación.",
            "roles": "Todos"
        }
    ]

    for pagina in paginas:
        doc.add_paragraph(f"**{pagina['ruta']} - {pagina['nombre']}**")
        doc.add_paragraph(f"Descripción: {pagina['descripcion']}")
        doc.add_paragraph(f"Roles autorizados: {pagina['roles']}")
        doc.add_paragraph("")  # Espacio

def agregar_especificacion_bd(doc):
    """Agrega especificación completa de la base de datos"""
    doc.add_heading('Especificación Completa de Base de Datos', level=2)

    tablas = [
        {
            "nombre": "usuarios",
            "descripcion": "Gestión de autenticación y control de acceso",
            "campos": [
                "id (PK, SERIAL)",
                "nombre (VARCHAR(255), NOT NULL)",
                "correo (VARCHAR(255), UNIQUE, NOT NULL)",
                "username (VARCHAR(100), UNIQUE)",
                "password (TEXT, NOT NULL) - Encriptado bcrypt",
                "rol_id (FK → roles.id)",
                "empleado_id (FK → empleados.id, NULL)",
                "created_at (TIMESTAMP, DEFAULT NOW())"
            ]
        },
        {
            "nombre": "empleados",
            "descripcion": "Información completa de empleados activos",
            "campos": [
                "id (PK, SERIAL)",
                "nombre (VARCHAR(255), NOT NULL)",
                "apellido (VARCHAR(255))",
                "documento (VARCHAR(20), UNIQUE, NOT NULL)",
                "correo (VARCHAR(255))",
                "telefono (VARCHAR(20))",
                "salario (DECIMAL(10,2))",
                "fecha_ingreso (DATE, NOT NULL)",
                "fecha_nacimiento (DATE)",
                "cargo (VARCHAR(255))",
                "departamento_id (FK → departamentos.id)",
                "estado (VARCHAR(50), DEFAULT 'activo')",
                "foto_url (TEXT) - URL Supabase Storage",
                "created_at (TIMESTAMP, DEFAULT NOW())",
                "updated_at (TIMESTAMP, DEFAULT NOW())"
            ]
        },
        {
            "nombre": "exempleados",
            "descripcion": "Historial de empleados retirados para auditoría",
            "campos": [
                "id (PK, SERIAL)",
                "nombre (VARCHAR(255), NOT NULL)",
                "apellido (VARCHAR(255))",
                "documento (VARCHAR(20))",
                "fecha_salida (DATE, NOT NULL)",
                "motivo_salida (TEXT)",
                "created_at (TIMESTAMP, DEFAULT NOW())"
            ]
        },
        {
            "nombre": "reportes",
            "descripcion": "Reportes administrativos de empleados",
            "campos": [
                "id (PK, SERIAL)",
                "empleado_id (FK → empleados.id)",
                "descripcion (TEXT, NOT NULL)",
                "fecha (DATE, DEFAULT NOW())",
                "creado_por (FK → usuarios.id)",
                "estado (VARCHAR(50), DEFAULT 'pendiente')",
                "created_at (TIMESTAMP, DEFAULT NOW())"
            ]
        },
        {
            "nombre": "documentos_empleado",
            "descripcion": "Documentos asociados a empleados",
            "campos": [
                "id (PK, SERIAL)",
                "empleado_id (FK → empleados.id)",
                "nombre_original (VARCHAR(255), NOT NULL)",
                "tipo (VARCHAR(100)) - MIME type",
                "url (TEXT, NOT NULL) - URL Supabase Storage",
                "created_at (TIMESTAMP, DEFAULT NOW())"
            ]
        },
        {
            "nombre": "certificados",
            "descripcion": "Registro de certificados generados",
            "campos": [
                "id (PK, SERIAL)",
                "empleado_id (FK → empleados.id)",
                "tipo_certificado (VARCHAR(100), DEFAULT 'laboral')",
                "fecha_generacion (DATE, DEFAULT NOW())",
                "created_at (TIMESTAMP, DEFAULT NOW())"
            ]
        },
        {
            "nombre": "departamentos",
            "descripcion": "Estructura organizacional de la empresa",
            "campos": [
                "id (PK, SERIAL)",
                "nombre (VARCHAR(255), UNIQUE, NOT NULL)",
                "created_at (TIMESTAMP, DEFAULT NOW())"
            ]
        },
        {
            "nombre": "roles",
            "descripcion": "Definición de roles del sistema",
            "campos": [
                "id (PK, SERIAL)",
                "nombre (VARCHAR(50), UNIQUE, NOT NULL)",
                "created_at (TIMESTAMP, DEFAULT NOW())"
            ]
        },
        {
            "nombre": "estados_empleados",
            "descripcion": "Estados posibles de los empleados",
            "campos": [
                "id (PK, SERIAL)",
                "nombre (VARCHAR(50), UNIQUE, NOT NULL)",
                "created_at (TIMESTAMP, DEFAULT NOW())"
            ]
        }
    ]

    for tabla in tablas:
        doc.add_paragraph(f"**Tabla: {tabla['nombre']}**")
        doc.add_paragraph(f"Descripción: {tabla['descripcion']}")
        doc.add_paragraph("Campos:")
        for campo in tabla['campos']:
            doc.add_paragraph(f"  - {campo}")
        doc.add_paragraph("")  # Espacio

def agregar_conclusiones_recomendaciones(doc):
    """Agrega conclusiones y recomendaciones"""
    doc.add_heading('Conclusiones y Recomendaciones', level=1)

    doc.add_heading('Conclusiones', level=2)
    doc.add_paragraph("""
El desarrollo de la plataforma web de gestión de talento humano para Formacero S.A.S. ha sido completado exitosamente, cumpliendo con todos los objetivos planteados inicialmente. El sistema implementado ofrece una solución integral que digitaliza y automatiza los procesos de gestión de recursos humanos, mejorando significativamente la eficiencia operativa del área administrativa.

**Logros principales:**
- Implementación completa del CRUD de empleados con gestión documental
- Sistema de autenticación robusto con JWT y control de roles
- Generación automática de certificados laborales en formato PDF
- Interfaz responsive adaptable a dispositivos móviles y escritorio
- Integración con Supabase para persistencia de datos y archivos
- Cumplimiento de estándares de seguridad y protección de datos

**Arquitectura implementada:**
La solución sigue una arquitectura cliente-servidor desacoplada, utilizando React para el frontend y Node.js con Express para el backend. La base de datos PostgreSQL gestionada por Supabase garantiza integridad y escalabilidad. La implementación de middlewares de autenticación y autorización asegura el control de acceso según roles.

**Calidad del software:**
El sistema ha sido desarrollado siguiendo buenas prácticas de programación, incluyendo separación de responsabilidades, reutilización de componentes y documentación técnica. Las pruebas realizadas confirman el correcto funcionamiento de todas las funcionalidades requeridas.""")

    doc.add_heading('Recomendaciones', level=2)
    doc.add_paragraph("""
**Para el mantenimiento y evolución del sistema:**

1. **Monitoreo continuo**: Implementar herramientas de monitoreo de rendimiento y errores para detectar problemas proactively.

2. **Actualizaciones de seguridad**: Mantener actualizadas las dependencias de Node.js, React y Supabase para corregir vulnerabilidades.

3. **Backup y recuperación**: Continuar con los backups automáticos diarios y probar regularmente los procedimientos de restauración.

4. **Capacitación**: Proporcionar capacitación continua al personal de RRHH sobre el uso del sistema.

**Mejoras futuras sugeridas:**

1. **Módulo de nómina**: Aunque inicialmente no se planteó, sería beneficioso integrar un módulo de cálculo de nómina y aportes legales.

2. **Módulo de vacaciones**: Sistema para gestión de solicitudes y control de días de vacaciones.

3. **Integración biométrica**: Posibilidad de integrar con dispositivos biométricos para control de asistencia.

4. **API externa**: Exposición de APIs para integración con otros sistemas empresariales.

5. **Analíticas avanzadas**: Dashboard con métricas de RRHH y reportes ejecutivos.

6. **Firma digital**: Implementación de certificados laborales con firma digital notarial.

**Recomendaciones técnicas:**

1. **Escalabilidad**: La arquitectura actual permite escalar horizontalmente agregando más instancias del servidor.

2. **Performance**: Implementar caché en Redis para mejorar tiempos de respuesta en consultas frecuentes.

3. **Testing**: Establecer un proceso de integración continua con pruebas automatizadas.

4. **Documentación**: Mantener actualizada la documentación técnica y de usuario.""")

def agregar_bibliografia(doc):
    """Agrega bibliografía"""
    doc.add_heading('Bibliografía', level=1)

    referencias = [
        "IEEE. (1998). IEEE Recommended Practice for Software Requirements Specifications. IEEE Std 830-1998.",
        "ISO/IEC. (2017). ISO/IEC 12207:2017 Systems and software engineering — Software life cycle processes.",
        "W3C. (2018). Web Content Accessibility Guidelines (WCAG) 2.1. https://www.w3.org/TR/WCAG21/",
        "Congreso de la República de Colombia. (2012). Ley 1581 de 2012 - Protección de Datos Personales.",
        "React Documentation. (2024). https://react.dev/",
        "Node.js Documentation. (2024). https://nodejs.org/docs/",
        "Supabase Documentation. (2024). https://supabase.com/docs",
        "Express.js Documentation. (2024). https://expressjs.com/",
        "JWT.io. (2024). JSON Web Tokens. https://jwt.io/",
        "Mozilla Developer Network. (2024). Web APIs. https://developer.mozilla.org/en-US/docs/Web/API"
    ]

    for i, ref in enumerate(referencias, 1):
        doc.add_paragraph(f"[{i}] {ref}")

def agregar_manual_usuario(doc):
    """Agrega manual de usuario básico"""
    doc.add_heading('Manual de Usuario', level=1)

    doc.add_heading('Instalación y Configuración', level=2)
    doc.add_paragraph("""
**Requisitos del sistema:**
- Navegador web moderno (Chrome 90+, Firefox 88+, Edge 90+, Safari 14+)
- Conexión a internet estable
- Resolución de pantalla mínima: 1024x768

**Configuración inicial:**
1. El sistema está alojado en servidor dedicado
2. No requiere instalación local por parte del usuario
3. Acceso mediante URL proporcionada por el administrador

**Variables de entorno configuradas:**
- SUPABASE_URL: URL de conexión a base de datos
- SUPABASE_ANON_KEY: Clave de acceso anónimo
- EMAIL_HOST/PORT/USER/PASS: Configuración SMTP
- FRONTEND_URL: URL de la aplicación frontend""")

    doc.add_heading('Guía de Uso por Roles', level=2)

    # Administrador
    doc.add_paragraph("**Rol: Administrador**")
    doc.add_paragraph("""
**Funciones principales:**
1. **Gestión de Empleados**: Crear, editar, consultar y retirar empleados
2. **Gestión Documental**: Subir y gestionar documentos de empleados
3. **Generación de Certificados**: Crear certificados laborales en PDF
4. **Reportes**: Crear y gestionar reportes administrativos
5. **Auditoría**: Consultar logs de cambios en el sistema

**Flujo de trabajo típico:**
1. Iniciar sesión con credenciales de administrador
2. Acceder al dashboard principal
3. Registrar nuevos empleados usando el formulario completo
4. Gestionar información existente mediante la lista de empleados
5. Generar certificados cuando sean solicitados""")

    # Empleado
    doc.add_paragraph("**Rol: Empleado**")
    doc.add_paragraph("""
**Funciones disponibles:**
1. **Consulta de Perfil**: Ver información personal y laboral
2. **Edición Limitada**: Actualizar correo, teléfono y foto de perfil
3. **Gestión Documental**: Descargar documentos personales
4. **Generación de Certificados**: Solicitar certificado laboral

**Flujo de trabajo típico:**
1. Iniciar sesión con credenciales personales
2. Revisar información en el perfil
3. Actualizar datos de contacto si es necesario
4. Descargar certificado laboral cuando se requiera""")

    doc.add_heading('Solución de Problemas', level=2)
    doc.add_paragraph("""
**Problemas comunes y soluciones:**

1. **No puedo iniciar sesión:**
   - Verificar que usuario y contraseña sean correctos
   - Usar la función "Recuperar contraseña" si es necesario
   - Contactar al administrador si el problema persiste

2. **Error al subir archivos:**
   - Verificar que el archivo no exceda 2MB
   - Asegurarse de que sea formato JPG, PNG o PDF
   - Comprobar conexión a internet

3. **Página no carga correctamente:**
   - Actualizar el navegador
   - Limpiar caché del navegador
   - Intentar con otro navegador compatible

4. **Certificado no se genera:**
   - Verificar que todos los datos del perfil estén completos
   - Contactar al administrador para verificar permisos

**Soporte técnico:**
Para soporte técnico, contactar al administrador del sistema o al equipo de desarrollo.""")

def actualizar_alcance(doc):
    """Actualiza la sección de alcance para excluir nómina y vacaciones"""
    # Buscar la sección de alcance
    alcance_encontrado = False
    for para in doc.paragraphs:
        if "Alcance de la solución:" in para.text:
            alcance_encontrado = True
            continue
        if alcance_encontrado and ("Para garantizar la viabilidad" in para.text):
            para.text = """Para garantizar la viabilidad del proyecto dentro del cronograma establecido, no se incluyen las siguientes funcionalidades que fueron consideradas fuera del alcance inicial:

• Liquidación de nómina y cálculos de pagos o aportes legales
• Módulo de reclutamiento externo
• Integración con dispositivos biométricos de control de asistencia
• Implementación de firmas digitales notariadas con validez legal externa
• Módulo de gestión de vacaciones (solicitudes, aprobaciones, control de días)
• Módulo de gestión de incapacidades y licencias

Estas funcionalidades quedan como propuestas para futuras versiones del sistema."""
            break

if __name__ == "__main__":
    actualizar_documento_word()